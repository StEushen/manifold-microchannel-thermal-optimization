# -*- coding: utf-8 -*-
"""Build the revised competition paper as a .docx."""
import json, os
from PIL import Image
import numpy as np
import pandas as pd
from sklearn.linear_model import Ridge
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG = os.path.join(ROOT, "figs")
EQ = os.path.join(ROOT, "eqs")
OUT = os.path.join(ROOT, "output", "歧管式微通道芯片热管理系统_修订稿.docx")
os.makedirs(os.path.dirname(OUT), exist_ok=True)

rep = json.load(open(os.path.join(ROOT, "report.json"), encoding="utf-8"))

# ---------------------------------------------------------------------------
# refit models for the appendix coefficient table
# ---------------------------------------------------------------------------
DATA = os.path.join(ROOT, "data", "附件2_数据.xlsx")
df = pd.read_excel(DATA, header=1)
df.columns = ["no", "beta", "lam", "N", "R", "dP", "UT"]
fin = df[df.beta > 0].reset_index(drop=True)
Xf = fin[["beta", "lam", "N"]].values.astype(float)
Yf = fin[["R", "dP", "UT"]].values.astype(float)
MU = np.array([0.1875, 3.75, 6.0]); SD = np.array([0.073951, 0.559017, 2.828427])

def design(Zm, deg=3):
    cols = [np.ones(len(Zm))]
    for d in range(1, deg + 1):
        for i in range(d + 1):
            for j in range(d - i + 1):
                l = d - i - j
                cols.append((Zm[:, 0] ** i) * (Zm[:, 1] ** j) * (Zm[:, 2] ** l))
    return np.column_stack(cols)

Ph = design((Xf - MU) / SD)
fit_models = [Ridge(alpha=1e-10).fit(Ph, Yf[:, k]) for k in range(3)]
coefs = []
for m in fit_models:
    c = m.coef_.copy()
    c[0] = c[0] + m.intercept_
    coefs.append(c)
TERMS = ["1", "zβ", "zλ", "zN", "zβ²", "zβzλ", "zβzN", "zλ²", "zλzN", "zN²",
         "zβ³", "zβ²zλ", "zβ²zN", "zβzλ²", "zβzλzN", "zβzN²", "zλ³", "zλ²zN", "zλzN²", "zN³"]

# ---------------------------------------------------------------------------
# document helpers
# ---------------------------------------------------------------------------
doc = Document()

sec = doc.sections[0]
sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
sec.left_margin = Cm(2.8); sec.right_margin = Cm(2.8)
USABLE_CM = 21.0 - 2.8 - 2.8

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.paragraph_format.line_spacing = 1.3
normal.paragraph_format.space_after = Pt(0)

def set_run(run, text=None, east="宋体", ascii_f="Times New Roman", size=10.5,
            bold=False, italic=False):
    if text is not None:
        run.text = text
    run.font.name = ascii_f
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def para(text, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=10.5,
         east="宋体", bold=False, space_before=0, space_after=0, line=1.3,
         left=0, first=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = align
    pf.line_spacing = line
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if left:
        pf.left_indent = Pt(left)
    if indent and first is None:
        pf.first_line_indent = Pt(size * 2)
    elif first is not None:
        pf.first_line_indent = Pt(first)
    r = p.add_run(text)
    set_run(r, east=east, size=size, bold=bold)
    return p

def h1(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(12); pf.space_after = Pt(6); pf.line_spacing = 1.3
    r = p.add_run(text)
    set_run(r, east="黑体", size=14, bold=True)
    pPr = p._p.get_or_add_pPr()
    ol = OxmlElement("w:outlineLvl"); ol.set(qn("w:val"), "0"); pPr.append(ol)
    return p

def h2(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(8); pf.space_after = Pt(4); pf.line_spacing = 1.3
    r = p.add_run(text)
    set_run(r, east="黑体", size=12, bold=True)
    pPr = p._p.get_or_add_pPr()
    ol = OxmlElement("w:outlineLvl"); ol.set(qn("w:val"), "1"); pPr.append(ol)
    return p

def add_equation(num):
    img = os.path.join(EQ, f"eq{num}.png")
    im = Image.open(img)
    wp, hp = im.size
    w_pt = wp / 300.0 * 72.0
    h_pt = hp / 300.0 * 72.0
    max_w = USABLE_CM * 0.86 * 28.3465
    max_h = 26.0
    scale = min(1.0, max_w / w_pt, max_h / h_pt)
    w_pt *= scale; h_pt *= scale
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(3); pf.space_after = Pt(3)
    pf.tab_stops.add_tab_stop(Cm(USABLE_CM / 2), WD_TAB_ALIGNMENT.CENTER)
    pf.tab_stops.add_tab_stop(Cm(USABLE_CM), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run("\t")
    set_run(r1, size=10.5)
    rp = p.add_run()
    rp.add_picture(img, width=Pt(w_pt))
    r2 = p.add_run(f"\t（{num}）")
    set_run(r2, size=10.5)
    return p

def add_figure(img, caption, width_cm=14.0):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.add_run().add_picture(os.path.join(FIG, img), width=Cm(width_cm))
    c = doc.add_paragraph()
    c.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(6)
    r = c.add_run(caption)
    set_run(r, size=9)

def set_cell(cell, text, bold=False, size=9, align="center"):
    cell.vertical_alignment = 1
    p = cell.paragraphs[0]
    p.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER,
                   "left": WD_ALIGN_PARAGRAPH.LEFT,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run(r, size=size, bold=bold)

def table_borders(table, header=True):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge, sz in [("top", 12), ("bottom", 12), ("left", 0),
                     ("right", 0), ("insideH", 0), ("insideV", 0)]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single" if sz else "none")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)
    if header:
        for cell in table.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            bd = OxmlElement("w:tcBorders")
            b = OxmlElement("w:bottom")
            b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6"); b.set(qn("w:color"), "000000")
            bd.append(b); tcPr.append(bd)

def add_table(headers, rows, widths=None, caption=None, size=9, first_left=False):
    if caption:
        c = doc.add_paragraph()
        c.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_before = Pt(6)
        r = c.add_run(caption)
        set_run(r, size=9, bold=False)
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for j, htxt in enumerate(headers):
        set_cell(t.rows[0].cells[j], htxt, bold=True, size=size)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            al = "left" if (first_left and j == 0) else "center"
            set_cell(t.rows[i].cells[j], str(val), size=size, align=al)
    if widths:
        for j, w in enumerate(widths):
            for i in range(len(rows) + 1):
                t.rows[i].cells[j].width = Cm(w)
    table_borders(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(it); run._r.append(f2)

# ---------------------------------------------------------------------------
# title & abstract
# ---------------------------------------------------------------------------
tp = doc.add_paragraph()
tp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(6); tp.paragraph_format.space_after = Pt(12)
r = tp.add_run("歧管式微通道芯片热管理系统的机理—数据融合多目标优化")
set_run(r, east="黑体", size=16, bold=True)

hp = doc.add_paragraph()
hp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
hp.paragraph_format.space_after = Pt(6)
r = hp.add_run("摘  要")
set_run(r, east="黑体", size=14, bold=True)

para("针对高热流密度芯片歧管式微通道热管理系统，本文研究针肋宽度比、歧管深高比和针肋排数对无量纲热阻、压降及温度非均匀性的耦合影响，并在偏好未知和制造误差存在时给出可解释的结构方案。首先，从稳态不可压缩 Navier–Stokes 方程、流固共轭传热方程出发，构建“歧管压力—支路流量—针肋间隙速度—局部换热—芯片温度”的水力与热力网络。模型表明，增大针肋宽度比或排数可增加湿润面积、强化扰动，但会缩小流通面积并累积局部阻力；增大歧管深高比可显著降低歧管流速和压力损失，却可能削弱射流冲击。因此三类参数对热学指标存在非单调作用，而对压降具有较明确的方向性。")
para("其次，识别附件 84 组数据中“4 个无针肋基准点 + 80 个带针肋完整因子点”的拓扑差异，建立条件代理模型：无针肋域采用一维二次基准曲线，带针肋域比较二次响应面、三次响应面和高斯过程。5 次重复的 5 折交叉验证表明，三次响应面在三项指标上的范围归一化 RMSE 分别为 0.000666%、0.459% 和 0.420%，兼具精度、光滑性与显式可解释性。随后，在 0.10 ≤ β ≤ 0.30、3 ≤ λ ≤ 4.5、N ∈ {2, 4, 6, 8, 10} 内实施约 15 万点的确定性网格搜索。以样本极差归一化后的三指标等权加权，得到综合推荐方案 (β, λ, N) = (0.220, 4.50, 4)，其预测性能为 (R*, Δp*, UT*) = (0.741183, 0.084698, 0.771878)。为处理应用偏好变化，进一步在权重单纯形上建立最小最大后悔模型，得到偏好鲁棒方案 (0.226, 4.50, 6)，其最大后悔值为 0.270865。最后，在明确声明的 β ± 0.01、λ ± 0.05 有界制造误差下进行 100000 次蒙特卡洛传播。两方案三项指标的变异系数均不超过 1.34%，但偏好鲁棒方案的压降对针肋宽度更敏感。本文最终建议：偏好明确且三指标等重要时采用 4 排等权方案；应用偏好尚未确定、需控制最坏决策损失时采用 6 排偏好鲁棒方案。所有结论仅适用于附件给定结构域和固定工况，不对流量、热负荷或区间外结构作无证据外推。")
kp = doc.add_paragraph()
kp.paragraph_format.space_after = Pt(6)
r = kp.add_run("关键词：共轭传热；歧管式微通道；响应面；多目标优化；最小最大后悔；蒙特卡洛")
set_run(r, size=10.5, bold=True)

# ---------------------------------------------------------------------------
# 一、问题重述
# ---------------------------------------------------------------------------
h1("一、问题重述")
h2("1.1  问题背景")
para("随着高功率密度芯片与三维集成封装技术的飞速发展[1-3]，芯片单位面积发热量显著剧增，若热量无法及时导出，局部“热点”将导致结温过高、芯片性能衰减乃至热应力失效[2]。传统直通式微通道虽然换热能力强[4]，但存在沿程压降大、出口温升高及温度分布不均的瓶颈[5-6]；为此，通过分配层将长流道截断为多个短流道的歧管式微通道（MMC）[5,8]结合针肋强化结构[7]，能够有效降低流动阻力并增强流体扰动。然而，高功率芯片的热管理设计不能单方面追求最低结温，针肋强化换热会增加局部阻力，过深的歧管虽有利于降低压降却可能削弱微通道入口处的冲击与混合，且较低的温度平均水平并不保证局部热点和热应力足够小。因此，针肋宽度比、歧管深高比与针肋排数等参数对热阻、压降与温度均匀性存在着复杂的非线性耦合与多目标制约关系，使本问题实质上成为一个受传热—流动耦合约束的小样本、多目标、混合离散设计优化问题。")
h2("1.2  问题重述")
para("针对附件 1 给出的歧管式微通道芯片热管理结构以及附件 2 的无量纲性能数据，本文需解决以下五个核心任务：")
for t in [
    "问题 1（机理建模与指标合理性）：从传热学与流体力学基本机理出发，建立结构参数与性能指标间的数学关系，分析针肋宽度比 β、歧管深高比 λ 和针肋排数 N 对无量纲热阻 R*th、压降 Δp* 及温度非均匀性 UT* 的影响规律，并阐明三指标作为综合评价依据的合理性。",
    "问题 2（代理模型构建与评估）：基于附件 2 中的 84 组样本数据，区分拓扑结构差异，构建高精度的代理模型以精准映射结构参数与各性能指标间的关系，并对其拟合与泛化预测性能进行定量评估。",
    "问题 3（多目标优化与综合方案）：以热阻、压降和温度非均匀性最小化为目标，建立多目标优化模型；基于合理的归一化与等权准则，在设计区间内搜索并确定综合最优结构参数方案。",
    "问题 4（偏好敏感性与鲁棒优化）：分析不同应用场景下指标权重变化对最优解的影响；引入最小最大后悔模型，求解对偏好变化不敏感、适应性强的偏好鲁棒设计方案。",
    "问题 5（加工误差与工况波动分析）：考查几何参数公差及运行工况波动等实际因素，建立蒙特卡洛传播模型，评估最优方案对小范围扰动的敏感性与工程稳定性。",
]:
    para(t)

# ---------------------------------------------------------------------------
# 二、问题分析
# ---------------------------------------------------------------------------
h1("二、问题分析")
h2("2.1  问题一的分析")
para("问题一要求从物理机理层面建立结构与性能的桥梁。首先，利用连续介质假设与能量守恒定律确定芯片总发热量；其次，建立连续流固共轭传热方程以及降阶的水力-热力网络模型；最后，推导针肋阻塞效应、湿润面积变化以及歧管分流对三项指标的作用机制，证明热阻与非均匀性的非单调性以及压降的单调方向性，从而确立三指标分别衡量“最高结温、泵功能耗、热应力风险”的互相不可替代的合理性。")
h2("2.2  问题二的分析")
para("问题二要求基于 84 组样本构建数据驱动代理模型。分析数据可知，样本包含 4 个无针肋点（β = 0、N = 0）与 80 个带针肋完整因子点（4 × 4 × 5）。由于有无针肋属于物理拓扑突变，强行统一插值会导致伪响应。因此，采取“条件代理模型”策略：无针肋采用一维二次拟合，带针肋对比二次/三次响应面与高斯过程（GPR）。通过重复 5 折交叉验证（CV）挑选最优模型。")
h2("2.3  问题三的分析")
para("问题三属于多目标结构优化。利用问题二训练出的高精度三次响应面代理模型，在限定的设计空间（β ∈ [0.1, 0.3]、λ ∈ [3, 4.5]、N ∈ {2, 4, 6, 8, 10}）内，使用高分辨率网格搜索求解 Pareto 最优解集。将三项指标采用极差归一化后，以等权加权和与理想点距离两种准则，确定出兼顾散热、能耗与均匀性的综合推荐方案。")
h2("2.4  问题四的分析")
para("问题四分析决策偏好不确定性。在实际工程中，不同应用（如高算力爆发场景偏好热阻，低功耗场景偏好压降）对权重的需求不同。首先通过在权重单纯形上扫描，分析最优解的迁移轨迹；随后，引入决策论中的“最小最大后悔（Minimax Regret）”准则，寻找在任意可能的权重偏好下，最大性能损失最小的鲁棒设计方案。")
h2("2.5  问题五的分析")
para("问题五关注实际制造与运行波动下的稳定性。对连续变量 β、λ 施加无偏有界制造误差（β ± 0.01、λ ± 0.05），采用蒙特卡洛抽样传播，统计性能指标的均值、标准差与变异系数（CV）；对离散排数 N 进行邻域扰动分析，评估推荐方案在公差扰动下的几何与水热稳定性。")

# ---------------------------------------------------------------------------
# 三、模型假设
# ---------------------------------------------------------------------------
h1("三、模型假设")
for t in [
    "1. 系统达到稳态，去离子水为单相、不可压缩牛顿流体；在题设温度区间内物性取常数。",
    "2. 所有固体壁面无滑移；忽略辐射；流固界面不存在额外接触热阻，温度与法向热流连续。",
    "3. 芯片热源均匀，外表面保留题设 hair = 10 W/(m²·K) 的自然对流边界。",
    "4. 不预先指定层流或湍流。摩擦因子与 Nusselt 数写成 Reynolds 数和几何参数的函数，由样本响应闭合。",
    "5. 题设“入口质量流量 1 g/s”解释为三个入口合计质量流量。若原始仿真实际对每个入口分别施加 1 g/s，只需在能量方程中将总流量替换为 3 g/s；附件内所有样本工况一致，因此结构—无量纲性能代理映射不受这一文字口径影响。",
    "6. 附件没有给出完整通道截面、针肋节距及三个输出的归一化基准。本文不虚构这些量；机理模型保留符号，数值预测直接学习题给无量纲输出。",
]:
    para(t)

# ---------------------------------------------------------------------------
# 四、符号说明
# ---------------------------------------------------------------------------
h1("四、符号说明")
add_table(
    ["符号", "含义", "符号", "含义"],
    [
        ["dp, wc", "针肋直径、通道宽度", "Rth", "热阻"],
        ["ṁi", "第 i 支路质量流量", "UT", "温度非均匀性"],
        ["β, λ, N", "宽度比、深高比、排数", "f, K", "沿程摩擦因子、局部损失系数"],
        ["Hm, Hc", "歧管层、微通道层高度", "Cm", "支路流量变异系数"],
        ["Af, Pw, Dh", "流通面积、润湿周长、水力直径", "Δp", "系统压降"],
        ["Re, Pr, Nu", "雷诺数、普朗特数、努塞尔数", "Q", "芯片总发热功率"],
        ["h", "对流换热系数", "ṁΣ", "入口总质量流量"],
        ["wk", "第 k 个指标权重", "ε, ϕ", "孔隙率、针肋平面占据率"],
    ],
    widths=[3.3, 4.5, 3.3, 4.5], size=9,
)

# ---------------------------------------------------------------------------
# 五、模型的建立与求解
# ---------------------------------------------------------------------------
h1("五、模型的建立与求解")

h2("5.1  传热—流动机理模型")
h2("5.1.1  结构与总能量尺度")
para("系统结构及热流、水流路径如图 2。水由三个顶部入口进入歧管层，经分配后穿入微通道换热层，绕过针肋并由两个侧部出口流出；芯片热量经氮化铝层传至微通道壁和针肋表面。芯片散热区域在平面上为 10 mm × 10 mm 的方形区域，芯片热源位于上表面中心区域，可视为均匀体热源；在附件未给出热源具体平面尺寸的情况下，本文按热源区域 6 mm × 10 mm、厚度 200 µm 作量级估计，体热源强度为 q‴ = 5 × 10⁹ W/m³，故：")
add_figure("fig2_structure.png", "图 2  歧管式微通道热管理系统结构与热流、水流路径示意图", width_cm=15.2)
add_equation(1)
para("对整个装置作稳态能量平衡：")
add_equation(2)
para("其中 T̄out,m 为两个出口按质量流量加权的混合温度。在 Qair 相对较小时，若总流量为 1 g/s，则平均水温升量级为：")
add_equation(3)
para("式（3）说明：固定热源和固定总流量下，结构不能凭空改变水的总焓增；它真正改变的是各支路分到的水量、局部壁水温差、热点位置和维持流量所需的压力。附件 1 给定的冷却介质与结构材料物性列于表 1。")
add_table(
    ["材料", "密度 ρ/(kg·m⁻³)", "比热 cp/(J·kg⁻¹·K⁻¹)", "导热系数 k/(W·m⁻¹·K⁻¹)", "动力黏度 μ/(kg·m⁻¹·s⁻¹)"],
    [
        ["去离子水", "998.2", "4182", "0.6", "0.001003"],
        ["氮化铝", "3260", "700", "200", "—"],
        ["芯片热源等效材料", "19320", "130", "298", "—"],
    ],
    widths=[3.8, 2.8, 3.0, 3.0, 3.0],
    caption="表 1  冷却介质与结构材料物性参数",
)

h2("5.1.2  共轭传热控制方程")
para("流体域的质量、动量和能量守恒为：")
add_equation(4)
add_equation(5)
add_equation(6)
para("芯片与氮化铝固体域满足：")
add_equation(7)
para("流固界面满足温度和法向热流连续：")
add_equation(8)
para("入口给定 Tin = 293 K 和质量流量，出口给定 pout = 0 Pa 表压，固体外侧满足：")
add_equation(9)
para("式（4）—（9）与由 (β, λ, N) 决定的流固区域共同构成精确边值问题，其数值求解可参照经典传热与流体力学理论[9]。题设工况与边界条件为：冷却液入口质量流量 1 g/s、入口温度 293 K、出口压力 0 Pa（表压）、外界环境温度 293 K、散热器外侧面对空气自然对流换热系数 10 W/(m²·K)、芯片体热源强度 5 × 10⁹ W/m³。下面将式（4）—（9）降阶为能够显示结构参数作用路径的水力与换热网络。")

h2("5.1.3  歧管—微通道水力网络")
para("将换热区沿歧管方向划分为 J 个单元，第 j 条微通道支路流量为 ṁj。供水歧管和集水歧管中轴向流量分别为 Mjs、Mjc，节点质量守恒为：")
add_equation(10)
para("更一般地，若把三个入口、两个出口及全部流道视为有向图，节点—边关联矩阵 B 满足：")
add_equation(11)
para("第 j 支路两端压差与流量匹配：")
add_equation(12)
para("水力直径和 Reynolds 数为：")
add_equation(13)
para("若一个歧管单元平面面积为 Apl，每排有 n⊥ 根圆柱针肋，则针肋平面占据率与等效孔隙率为：")
add_equation(14)
add_equation(15)
para("在支路流量给定时，针肋区速度尺度满足：")
add_equation(16)
para("因此 β 或 N 增大都会降低孔隙率、提高间隙速度；同时 Kp 随阻塞增强，排数还使损失近似逐排累加。")
para("歧管若近似为宽 Wm、高 Hm = λHc 的矩形流道，则：")
add_equation(17)
para("宽浅层流矩形歧管和局部惯性损失的尺度分别为：")
add_equation(18)
para("由此，在题给结构域、总流量固定且不发生旁路拓扑变化时，可判断：")
add_equation(19)
para("这里 ΔNΔp = Δp(N + 2) − Δp(N)。式（19）给的是有明确前提的物理趋势，不是脱离具体几何的无条件定理。")

h2("5.1.4  水侧换热网络与热阻")
para("第 j 支路的混合平均水温满足：")
add_equation(20)
para("若单元内壁温近似为常数，积分得：")
add_equation(21)
para("若针肋贯穿微通道高度，湿润面积可写为：")
add_equation(22)
para("其中 νb 表示需要扣除的有效受热投影面数，取决于针肋与一侧或两侧壁面的连接形式。对流换热系数为：")
add_equation(23)
para("以芯片最高温度定义热阻：")
add_equation(24)
para("其中：")
add_equation(25)
para("对 x ∈ {β, N} 求对数灵敏度：")
add_equation(26)
para("式（26）揭示了热阻非单调性的根源：针肋初期增加时，Nu·Awet 的增长占优，热阻下降；继续加粗或加密后，尾迹滞流、流量重分配与有效面积缩减占优，热阻收益饱和甚至反弹。歧管深度通过改变支路流量和射流速度间接作用于热阻，因此守恒律本身不能保证 Rth 对 λ 单调。")

h2("5.1.5  温度非均匀性")
para("附件没有公布无量纲温度非均匀性的原始基准。机理层以等面积采样温度的标准差型指标说明其来源：")
add_equation(27)
para("其中 Θref > 0 为固定温差基准。数值层直接学习附件给出的 UT*，不以式（27）擅自替换原始指标。")
para("定义支路流量变异系数：")
add_equation(28)
para("局部温升近似为：")
add_equation(29)
para("令 ṁi = m̄(1 + δi)、hi ∝ ṁiai，对式（29）作一阶展开可得：")
add_equation(30)
para("因此温度离散度的一阶量级与流量不均匀 Cm 正相关，并叠加固体扩散路径和针肋尾迹造成的局部差异。适度加深歧管有利于减小歧管压力梯度和 Cm，但过深又可能降低射流速度；适量针肋能均化换热，过粗或过多则会制造低速区。因此 UT 对三参数均可能存在中间最优。")

h2("5.1.6  三指标作为综合评价依据的合理性")
para("热阻决定最高温度、压降决定泵功、温度非均匀性关联受约束材料的热应力量级：")
add_equation(31)
add_equation(32)
add_equation(33)
para("故三指标分别衡量绝对温度水平、流动能耗与空间热应力。它们在数学上互不决定：相同最高温度可以对应不同温度场，相同温度分布也可以需要不同泵功。删去任一项，都可能把过热、高能耗或存在严重热点的结构误判为优良结构。")

h2("5.1.7  样本主效应对机理的验证")
para("为避免拓扑差异混淆，图 3 只统计 80 个带针肋完整因子样本，并对另外两个变量取平均。")
para("从 β = 0.10 增至 0.20，热阻和非均匀性分别下降约 1.98% 和 2.50%，压降增加约 7.07%；继续增至 0.30，压降骤增约 28.21%，热阻和非均匀性分别反弹约 0.91% 和 2.79%。这与“换热面积/扰动强化—堵塞/尾迹损失”的竞争一致。")
para("λ 从 3.0 增至 4.5 时，压降下降约 32.15%，验证了式（18）的尺度判断；热阻总体增加约 1.94%，均匀性在 λ = 3.5 附近最好，说明压力分配改善与冲击换热减弱同时存在。")
para("N 从 2 增至 10 时，热阻下降约 2.18%，但压降增加约 30.90%，且热阻边际收益递减；非均匀性在 4 排最小，随后因尾迹累积和支路重分配而恶化。主效应只能显示平均趋势，βλ、βN、λN 交互作用仍需代理模型定量识别。")
add_figure("fig3_maineffects.png", "图 3  结构参数对三项性能指标的主效应（80 个带针肋样本，其余两变量取平均）", width_cm=14.8)

h2("5.2  条件代理模型")
h2("5.2.1  数据结构与拓扑分组")
para("附件 2 共 84 个样本，其中：80 个带针肋点构成 4 × 4 × 5 完整因子网格，β ∈ {0.10, 0.15, 0.20, 0.30}、λ ∈ {3, 3.5, 4, 4.5}、N ∈ {2, 4, 6, 8, 10}；4 个无针肋点满足 β = 0、N = 0，仅改变参数 λ。有无针肋是实体拓扑改变，而不是普通连续变化。若强行用一个三变量光滑面连接全部样本，模型会在 β → 0、N > 0 等不存在的几何上产生伪预测。因此定义条件代理模型：")
add_equation(34)
para("无针肋组只有 4 点，仅用一维二次曲线表示基准趋势；所有推荐方案均位于带针肋完整数据域。")
h2("5.2.2  候选模型与标准化")
para("比较二次响应面、三次响应面和高斯过程回归。三次响应面写为：")
add_equation(35)
para("标准化变量为：")
add_equation(36)
para("标准化改善多项式设计矩阵的数值条件；三次项能够表达图 3 中的转折、曲率及二阶交互随第三变量变化的现象。为抑制极弱病态性，对系数施加 10⁻¹⁰ 的 Ridge 正则，其量级足够小，不改变响应面的实际拟合[10]。")
h2("5.2.3  交叉验证与模型选择")
para("在 80 个带针肋点上实施 5 次重复的 5 折交叉验证。评价指标为：")
add_equation(37)
para("图 4 给出一次确定性五折留出预测的样本—预测一致性，图 5 汇总重复交叉验证的候选模型误差。")
add_figure("fig4_cv_parity.png", "图 4  三次响应面一次确定性 5 折交叉验证的样本—预测一致性", width_cm=15.0)
add_figure("fig5_model_cmp.png", "图 5  三种候选模型重复 5 折交叉验证误差对比（5 次重复）", width_cm=13.5)
cv = rep["cubic"]
add_table(
    ["指标", "R²", "MAE", "RMSE", "范围归一化 RMSE"],
    [
        ["无量纲热阻", "1.000000", "2.910 × 10⁻⁷", "3.455 × 10⁻⁷", "0.000666%"],
        ["无量纲压降", "0.999555", "0.000246", "0.000585", "0.459%"],
        ["无量纲温度非均匀性", "0.999699", "0.000158", "0.000311", "0.420%"],
    ],
    widths=[3.6, 2.0, 2.2, 2.2, 2.6],
    caption="表 2  选定三次响应面的重复交叉验证性能",
)
para("三次响应面在三项输出上均优于二次响应面和高斯过程[10-11]，且连续可微、系数可导出，适合后续全域搜索、局部灵敏度及误差传播。热阻的近零误差说明该响应在给定网格内几乎完全符合低阶多项式结构；这不意味着模型在样本域外仍然精确，故后续禁止外推。")
h2("5.2.4  响应面解释")
para("固定 N = 4 的三项响应如图 6。热阻在较小 λ 和中等 β 一带较低；压降随 λ 增大明显降低，并在 β 过大时迅速上升；非均匀性在 β ≈ 0.22、较大 λ 附近形成低谷。三幅图的最优区域不重合，直观证明单目标最优不可能同时满足全部要求。")
add_figure("fig6_response.png", "图 6  固定 N = 4 时三项性能指标的响应面", width_cm=15.2)

h2("5.3  三目标优化与综合方案")
h2("5.3.1  优化模型")
para("针对三目标结构优化问题，采用加权和与 Pareto 支配相结合的方法求解[12]。建立三目标最小化模型：")
add_equation(38)
para("约束为：")
add_equation(39)
para("N 是离散可制造变量，不能取 5.37 排。对连续变量采用 Δβ = 0.001、Δλ = 0.01 的规则网格，并枚举五种排数，共评价约 15 万个带针肋方案。确定性网格搜索不依赖随机初值，分辨率足以支持 β 报告到三位小数、λ 报告到两位小数。")
h2("5.3.2  Pareto 支配与指标归一化")
para("若方案 xa 在三项目标上均不差于 xb，且至少一项严格更优，则称 xa 支配 xb。所有不被其他可行方案支配的点构成 Pareto 集。不同指标的数值跨度不同，先按题给全部样本的实际极值归一化：")
add_equation(40)
para("题目没有声明指标优先级，因此综合方案采用等权损失：")
add_equation(41)
para("没有采用熵权法：指标样本离散程度是数据性质，不等于使用方的价值偏好。等权表示“当前没有证据认为某一指标更重要”，并将在问题四中显式改变权重。图 7 展示支持型 Pareto 前沿。降低热阻通常伴随压降上升；低压降区需要较少针肋，却不能同时取得最低热阻；温度非均匀性最优区位于中间折中段。")
add_figure("fig7_pareto.png", "图 7  支持型 Pareto 前沿（0.05 权重网格扫描，115 个方案）", width_cm=13.2)
h2("5.3.3  优化结果")
para("等权综合推荐为：")
add_equation(42)
para("该方案与均匀性优先方案接近，说明温度均匀性低谷与综合折中区域重合；其压降接近低压区，同时避免 2 排方案较高的热阻。以到归一化理想点的加权欧氏距离进行独立交叉检查，得到 (0.224, 4.50, 6)，仍落在“β 约 0.22、λ 取上界、N 为 4—6”的同一区域，且与严格最优点的距离差异小于 10⁻⁵，说明结论不依赖唯一一种综合准则。")
add_table(
    ["决策准则", "β", "λ", "N", "热阻", "压降", "非均匀性"],
    [
        ["热阻优先 (0.8, 0.1, 0.1)", "0.228", "3.17", "10", "0.720555", "0.151744", "0.813938"],
        ["压降优先 (0.1, 0.8, 0.1)", "0.223", "4.50", "2", "0.759933", "0.076083", "0.779019"],
        ["均匀性优先 (0.1, 0.1, 0.8)", "0.224", "4.50", "4", "0.741114", "0.085117", "0.771783"],
        ["等权综合推荐", "0.220", "4.50", "4", "0.741183", "0.084698", "0.771878"],
        ["等权理想点距离", "0.224", "4.50", "6", "0.734381", "0.097795", "0.790194"],
    ],
    widths=[3.7, 1.1, 1.1, 1.0, 1.7, 1.7, 1.7],
    caption="表 3  不同决策准则下的参数配置与性能指标对比",
)

h2("5.4  偏好变化与偏好鲁棒设计")
h2("5.4.1  权重变化的影响")
para("令权重属于三维单纯形：")
add_equation(43)
para("图 8 将五类典型方案的归一化损失画在同一坐标上。热阻优先方案显著牺牲压降；压降优先方案牺牲热阻；4 排方案在非均匀性上占优；6 排理想点方案以中等压降换取更低热阻。权重改变不是数值扰动，而是评价标准改变，最优方案沿 Pareto 前沿迁移是合理现象。")
add_figure("fig8_weights.png", "图 8  不同应用偏好下五种典型方案的归一化加权损失", width_cm=15.0)
h2("5.4.2  最小最大后悔模型")
para("若应用偏好尚未确定，只比较平均加权得分会掩盖某些合法权重下的严重失配。定义方案 x 在权重 w 下的后悔值[13]：")
add_equation(44)
para("其中 P 为支持型 Pareto 候选集。选择最小最大后悔方案：")
add_equation(45)
para("计算时以 0.05 步长扫描权重以构造 115 个支持型 Pareto 方案，再以 0.01 步长在完整权重单纯形上评价后悔。结果为：")
add_equation(46)
para("性能为：")
add_equation(47)
para("最大后悔值为 0.270865。该数是归一化加权损失之差，不是百分比。图 9 中偏好鲁棒方案处于最大后悔值下包络。它相对等权 4 排方案降低热阻，但提高压降与非均匀性；其价值不是“三项性能都更好”，而是在热阻、泵功与均匀性偏好发生切换时控制最坏决策损失。严格地说，该结论在所枚举权重网格与支持型 Pareto 候选集上成立，不宣称对连续无限权重空间给出解析全局证明。")
add_figure("fig9_regret.png", "图 9  最小最大后悔分析：左为后悔曲线切片，右为全部候选方案的最大后悔值", width_cm=15.2)

h2("5.5  制造误差与运行波动")
h2("5.5.1  局部弹性系数")
para("对连续参数 xj ∈ {β, λ}，定义无量纲局部弹性：")
add_equation(48)
para("它表示参数相对变化 1% 时，指标近似变化 Ekj%。由三次响应面作中心差分计算。排数 N 是离散布局，不能用“半排”扰动，因此以相邻实验层级 N ± 2 作情景比较。")
h2("5.5.2  有界制造误差的蒙特卡洛传播")
para("附件没有给出实际加工公差。为避免把假设冒充测量值，显式设置情景：")
add_equation(49)
para("并将抽样严格截断在已验证设计域内。对等权综合方案和偏好鲁棒方案各进行 100000 次独立抽样，将样本代入代理模型，统计均值、标准差、P05 和 P95[14]。")
mc = rep["mc_eq"]; mr = rep["mc_rob"]
add_table(
    ["方案", "指标", "标称值", "均值", "标准差", "变异系数", "P95"],
    [
        ["等权综合", "热阻", "0.741183", "0.741535", "0.000412", "0.0556%", "0.742311"],
        ["等权综合", "压降", "0.084698", "0.084974", "0.000650", "0.7652%", "0.086078"],
        ["等权综合", "非均匀性", "0.771878", "0.772376", "0.000532", "0.0688%", "0.773347"],
        ["偏好鲁棒", "热阻", "0.734274", "0.734627", "0.000476", "0.0648%", "0.735512"],
        ["偏好鲁棒", "压降", "0.098231", "0.098527", "0.001313", "1.3322%", "0.100596"],
        ["偏好鲁棒", "非均匀性", "0.790195", "0.790619", "0.000393", "0.0497%", "0.791351"],
    ],
    widths=[2.0, 2.0, 1.9, 1.9, 1.8, 1.7, 1.7],
    caption="表 4  几何误差传播结果（β ± 0.01、λ ± 0.05，100000 次蒙特卡洛抽样）",
)
para("两方案在假设公差下三项指标变异系数均小于 1.34%，其中偏好鲁棒方案压降的变异系数最大（约 1.33%），说明局部几何稳定性较好。偏好鲁棒方案的压降变异系数更大，与其 β 对压降的弹性 0.510 一致；等权方案对应弹性为 0.259。两方案的 λ 压降弹性均为较大负值（分别为 −0.987 与 −0.864），说明继续加深歧管有利于降压，但推荐点已位于样本上边界，不能据此向 λ > 4.5 外推。")
h2("5.5.3  针肋排数的离散扰动")
nt = rep["N_table"]
add_table(
    ["标称设计 (β, λ)", "排数 N", "热阻", "压降", "温度非均匀性"],
    [
        ["等权参数 (0.220, 4.50)", "2", "0.759836", "0.076130", "0.779056"],
        ["等权参数 (0.220, 4.50)", "4", "0.741183", "0.084698", "0.771878"],
        ["等权参数 (0.220, 4.50)", "6", "0.734634", "0.096961", "0.790285"],
        ["鲁棒参数 (0.226, 4.50)", "4", "0.741097", "0.085343", "0.771784"],
        ["鲁棒参数 (0.226, 4.50)", "6", "0.734274", "0.098231", "0.790195"],
        ["鲁棒参数 (0.226, 4.50)", "8", "0.734383", "0.109646", "0.812385"],
    ],
    widths=[3.7, 1.5, 1.8, 1.8, 2.2],
    caption="表 5  不同标称设计及针肋排数下的性能指标对比",
)
para("从 4 排增至 6 排，热阻下降，但压降和非均匀性明显上升；从 6 排增至 8 排，热阻几乎不再改善，而水力和均匀性代价继续增大。因此 4—6 排是有效折中区，8 排以上缺乏综合优势。")
h2("5.5.4  工况波动的可识别边界")
para("附件所有样本均固定在同一质量流量、入口温度和热源条件，代理模型没有学习 ṁ、Tin、q‴ 三个工况维度。对压降只能由机理给出边界：黏性层流主导时 Δp ∝ ṁ，局部惯性主导时近似 Δp ∝ ṁ²。若流量增加 5%，则：")
add_equation(50)
para("热阻和温度非均匀性对工况的数值响应无法由现有 84 点可靠识别。实际工程扩展应对 ṁ、Tin、q‴ 各增加至少三个水平，将代理模型扩展为：")
add_equation(51)
para("在没有新增 CFD 或实验数据前，本文不报告伪精确的工况扰动结果。")

h2("5.6  结论")
for t in [
    "1. 结构参数通过“歧管压力分配—支路流量—针肋间隙速度—局部换热—温度场”链条影响性能。针肋宽度比和排数体现换热强化与流动阻塞的竞争；歧管深高比对压降显著有利，但对热阻和均匀性存在间接权衡。",
    "2. 带针肋域的三次响应面在重复交叉验证中达到 R² > 0.9995，三项范围归一化 RMSE 均小于 0.47%，能够支撑附件域内优化。无针肋结构因拓扑不同而单独处理。",
    "3. 等权综合推荐为 (0.220, 4.50, 4)，预测性能为 (0.741183, 0.084698, 0.771878)。理想点距离准则给出相邻的 6 排方案，交叉验证了 β ≈ 0.22、λ = 4.5、N = 4—6 的折中区。",
    "4. 权重未知时，最小最大后悔方案为 (0.226, 4.50, 6)，最大后悔值为 0.270865。它并非三指标逐项支配等权方案，而是在偏好切换下更稳健。",
    "5. 在 β ± 0.01、λ ± 0.05 的声明公差下，两方案三项指标变异系数均小于 1.34%；偏好鲁棒方案的压降对针肋宽度更敏感，证明偏好鲁棒与制造稳健需要分别评价。",
]:
    para(t)

# ---------------------------------------------------------------------------
# 六、模型的评价、改进与推广
# ---------------------------------------------------------------------------
h1("六、模型的评价、改进与推广")
h2("6.1  模型优点")
for t in [
    "• 物理与数据闭环。守恒方程、水力网络和 NTU 换热关系解释变量作用，响应面只承担附件缺失闭合关系的数值识别。",
    "• 尊重几何拓扑。无针肋基准与带针肋域条件建模，避免在不存在的 β = 0、N > 0 结构上插值。",
    "• 验证与优化分离。以重复交叉验证选择模型，再进行确定性全域网格搜索，降低过拟合和随机初值风险。",
    "• 区分两类鲁棒性。最小最大后悔处理评价偏好不确定性，蒙特卡洛处理制造误差；二者没有被混为同一个“鲁棒”概念。",
    "• 边界透明。排数保持离散，连续扰动不超出样本域，固定工况下不虚构流量或热负荷响应。",
]:
    para(t)
h2("6.2  模型局限")
for t in [
    "• 通道完整尺寸、针肋节距及原始无量纲定义缺失，无法以 CFD 独立复算每个样本。",
    "• 无针肋组只有 4 点，其基准曲线不具备强外推能力。",
    "• 推荐深高比位于 λ = 4.5 上边界，说明样本域可能截断了进一步降压的趋势；是否继续加深必须新增结构点并同时核查体积与加工约束。",
    "• 制造公差为情景假设，不是设备实测分布。进入工程设计后应以真实公差、相关性和工况概率替换均匀分布。",
    "• 支持型 Pareto 集由加权和生成，非凸前沿上的非支持解可能被遗漏；现有结论适用于扫描所得候选集。",
]:
    para(t)
h2("6.3  分层推荐")
para("常规综合设计：当三项指标没有明确优先级时，推荐：")
add_equation(52)
para("它具有最低附近的温度非均匀性和较低压降，在假设制造误差下压降变异系数为 0.765%。")
para("偏好未知设计：当产品应用尚未确定，未来可能在热阻、泵功和均匀性之间切换重点时，推荐：")
add_equation(53)
para("它牺牲部分压降和均匀性，以较低热阻控制权重变化下的最坏后悔；但应加强针肋宽度公差控制。")
para("进一步验证重点：在 λ = 4.5 附近加密 CFD/实验点，并新增 λ > 4.5 但满足体积约束的结构；围绕 β = 0.22、N = 4、6 进行流量和热负荷多水平试验；同时输出原始温度场和支路流量，以直接验证 Cm 与 UT 的关系。")

# ---------------------------------------------------------------------------
# 七、参考文献
# ---------------------------------------------------------------------------
h1("七、参考文献")
refs = [
    "Cang D, Dong Z, Lv S, et al. Design and intelligent optimization of TSV-based embedded microchannel heatsinks in 2.5D Packaging[J]. International Journal of Heat and Mass Transfer, 2026, 255: 127908.",
    "Zhang J, Sadiqbatcha S, Tan S X D. Hot-trim: Thermal and reliability management for commercial multicore processors considering workload dependent hot spots[J]. IEEE Transactions on Computer-Aided Design of Integrated Circuits and Systems, 2022, 42(7): 2290-2302.",
    "He W, Yin E, Zhou F, et al. Integrated manifold microchannels and near-junction cooling for enhanced thermal management in 3D heterogeneous packaging technology[J]. Energy, 2024, 305: 132263.",
    "Tuckerman D B, Pease R F W. High-performance heat sinking for VLSI[J]. IEEE Electron Device Letters, 1981, 2(5): 126-129.",
    "Harpole G M, Eninger J E. Micro-channel heat exchanger optimization[C]//Proceedings of the Seventh Annual IEEE Semiconductor Thermal Measurement and Management Symposium. Phoenix, AZ: IEEE, 1991: 59-63.",
    "Kandlikar S G, Grande W J. Evolution of microchannel flow passages—thermohydraulic performance and fabrication technology[J]. Heat Transfer Engineering, 2003, 24(1): 3-17.",
    "Peles Y, Koşar A, Mishra C, et al. Forced convective heat transfer across a pin fin micro heat sink[J]. International Journal of Heat and Mass Transfer, 2005, 48(17): 3615-3627.",
    "Drummond K P, Back D, Sinanis M D, et al. A hierarchical manifold microchannel heat sink array for high-heat-flux two-phase cooling of electronics[J]. International Journal of Heat and Mass Transfer, 2018, 117: 319-330.",
    "Bergman T L, Lavine A S, Incropera F P, et al. Fundamentals of Heat and Mass Transfer[M]. 7th ed. Hoboken: John Wiley & Sons, 2011.",
    "Box G E P, Draper N R. Response Surfaces, Mixtures, and Ridge Analyses[M]. 2nd ed. Hoboken: John Wiley & Sons, 2007.",
    "Rasmussen C E, Williams C K I. Gaussian Processes for Machine Learning[M]. Cambridge, MA: MIT Press, 2006.",
    "Deb K, Pratap A, Agarwal S, et al. A fast and elitist multiobjective genetic algorithm: NSGA-II[J]. IEEE Transactions on Evolutionary Computation, 2002, 6(2): 182-197.",
    "Savage L J. The theory of statistical decision[J]. Journal of the American Statistical Association, 1951, 46(253): 55-67.",
    "Metropolis N, Ulam S. The Monte Carlo method[J]. Journal of the American Statistical Association, 1949, 44(247): 335-341.",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    pf.left_indent = Pt(21)
    pf.first_line_indent = Pt(-21)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"[{i}] {ref}")
    set_run(r, size=9)

# ---------------------------------------------------------------------------
# 附录
# ---------------------------------------------------------------------------
h1("附录")
h2("附录 A：三次响应面回归系数")
para("表 6 给出带针肋域三次响应面（式（35））在三项指标上的回归系数。所有系数对应标准化变量 zβ、zλ、zN（式（36）），拟合采用 α = 10⁻¹⁰ 的 Ridge 正则，可据此直接复现正文全部预测结果。")
rows = []
for j, term in enumerate(TERMS):
    rows.append([term] + [f"{c[j]:.6e}" for c in coefs])
add_table(["项", "无量纲热阻 a(R*)", "无量纲压降 a(Δp*)", "无量纲温度非均匀性 a(UT*)"],
          rows, widths=[2.6, 4.0, 4.0, 4.0], caption="表 6  三次响应面回归系数（标准化变量）", size=8)
h2("附录 B：数值实验与复现说明")
for t in [
    "• 数据来源：附件 2 共 84 组样本，其中带针肋 80 组（β、λ、N 的 4 × 4 × 5 完整因子设计），无针肋 4 组（β = 0、N = 0）。",
    "• 代理模型：标准化三次响应面 + Ridge（α = 10⁻¹⁰）；5 次重复 5 折交叉验证（随机种子 6）；高斯过程采用 RBF 与 White 核组合。",
    "• 优化搜索：β 步长 0.001、λ 步长 0.01、N 枚举 5 个水平，共 151755 个带针肋方案；指标极值归一化采用全部 84 组样本；Pareto 候选由权重步长 0.05 的加权和扫描构造（115 个方案），后悔值在权重步长 0.01 的单纯形网格上评价。",
    "• 蒙特卡洛传播：每方案 100000 次独立抽样（随机种子 31），β、λ 分别取均匀分布并截断于设计域。",
    "• 软件环境：Python 3（NumPy、Pandas、scikit-learn、Matplotlib）；正文全部图表均可由上述设定直接复现。",
]:
    para(t, size=9.5, line=1.2)

add_page_number(sec)
doc.save(OUT)
print("saved:", OUT)
